import os
import pandas as pd
import re
from docx import Document
from config import users
from steps.user_login import user_login
from steps.manage_tabs import open_new_tab
from steps.driver_setup import setup_driver
from steps.navigate_and_select import navigate_and_select
from steps.select_document_type import select_document_type
from steps.input_case_number import input_case_number
from steps.upload_attachments import upload_attachments
from steps.submit_case import submit_case
from steps.pre_submit import pre_submit
from steps.verify_file import verify_file
import logging

# Configuración de logs
logging.basicConfig(
    filename='process_log.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)


def detectar_variables_en_escrito(docx_path):
    """
    Detecta todas las variables en el documento .docx que siguen el formato {{variable}}.
    """
    doc = Document(docx_path)
    texto_completo = "".join([para.text + "\n" for para in doc.paragraphs])
    pattern = r"\{\{(.*?)\}\}"
    return re.findall(pattern, texto_completo)


def obtener_valores_desde_excel(row, variables):
    """
    Obtiene los valores de las variables desde una fila específica de un DataFrame.
    """
    return {variable: row.get(variable, None) for variable in variables}


def reemplazar_variables_en_escrito(texto, valores_variables):
    """
    Reemplaza las variables en el texto del documento con los valores proporcionados.
    """
    for variable, valor in valores_variables.items():
        if valor is not None:
            pattern = r"\{\{\s*" + re.escape(variable) + r"\s*\}\}"
            texto = re.sub(pattern, str(valor), texto)
    return texto


def leer_documento_como_texto(docx_path):
    """
    Lee el contenido del archivo .docx como texto.
    """
    doc = Document(docx_path)
    return "".join([para.text + "\n" for para in doc.paragraphs])


def validar_campos_obligatorios(row, campos):
    """
    Verifica que los campos obligatorios estén presentes.
    """
    errores = []
    for campo in campos:
        if pd.isnull(row.get(campo)):
            errores.append(f"{campo} ausente")
    return errores


def execute_steps(driver, texto_modificado, id_sac, user_config, pdf_path, planilla_path):
    """
    Ejecuta los steps necesarios utilizando Selenium.
    """

    print("pdf_path", pdf_path)
    print("planilla_path", planilla_path)

    
    user_login(driver, user_config)
    open_new_tab(driver)
    navigate_and_select(driver)
    select_document_type(driver, texto_modificado)
    input_case_number(driver, id_sac)
    upload_attachments(driver, pdf_path, planilla_path)
    pre_submit(driver)
    submit_case(driver)


def main(excel_path, docx_path, user_key):
    if not verify_file(docx_path):
        logging.error("El archivo .docx no se encontró. Finalizando script.")
        return

    # Obtener configuración del usuario
    user_config = users.get(user_key.upper())
    if not user_config:
        logging.error(f"Usuario {user_key} no encontrado en la configuración.")
        return

    # Leer variables del documento
    variables_encontradas = detectar_variables_en_escrito(docx_path)
    logging.info(f"Variables encontradas en el documento: {variables_encontradas}")

    # Leer el archivo Excel
    df = pd.read_excel(excel_path)
    if df.empty:
        logging.error("El archivo Excel está vacío. Finalizando script.")
        return

    # Asegurar columna Observación
    if "Observación" not in df.columns:
        df["Observación"] = None

    for index, row in df.iterrows():
        if row.get("Observación") == "Título rectificado":
            logging.info(f"Fila {index}: El título ya fue rectificado. Saltando.")
            continue

        errores = validar_campos_obligatorios(row, ["IdSAC", "Planilla"])
        if errores:
            logging.warning(f"Fila {index}: {', '.join(errores)}. Saltando.")
            df.at[index, "Observación"] = ", ".join(errores)
            df.to_excel(excel_path, index=False)
            continue

        # Procesar valores
        id_sac = row["IdSAC"]
        valores_variables = obtener_valores_desde_excel(row, variables_encontradas)

        # Modificar el texto
        texto_completo = leer_documento_como_texto(docx_path)
        texto_modificado = reemplazar_variables_en_escrito(texto_completo, valores_variables)

        # Rutas de archivos
        base_dir = os.path.dirname(os.path.abspath(__file__))
        pdf_path = os.path.join(base_dir, "pdfs", f"{id_sac}_firmado.pdf")
        planilla_path = os.path.join(base_dir, "planillas", f"{row['Planilla']}.pdf")

        # Ejecutar pasos
        driver = setup_driver()
        try:
            execute_steps(driver, texto_modificado, id_sac, user_config, pdf_path, planilla_path)
            df.at[index, "Observación"] = "Título rectificado"
        except Exception as e:
            logging.error(f"Error en fila {index}: {e}")
            df.at[index, "Observación"] = f"Error: {e}"
        finally:
            driver.quit()

        df.to_excel(excel_path, index=False)

    logging.info("Procesamiento completado.")


if __name__ == "__main__":
    base_path = os.path.dirname(os.path.abspath(__file__))
    print(base_path)
    excel_path = os.path.join(base_path, "tables", "Causas 2021 para desistir - iniciadas dos veces.xlsx")
    docx_path = os.path.join(base_path, "certificates", "Escrito_base.docx")
    user_key = "FONSECA"
    main(excel_path, docx_path, user_key)
