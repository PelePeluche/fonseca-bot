import os
import pandas as pd
import re
from docx import Document
from config import users
from steps.user_login import user_login
from steps.manage_tabs import open_new_tab
from steps.driver_setup import setup_driver
from steps.navigate_and_select import navigate_and_select
from steps.select_document_type import select_document_type
from steps.input_case_number import input_case_number
from steps.skip_attachments import skip_attachments
from steps.submit_case import submit_case
from steps.pre_submit import pre_submit
from steps.verify_file import verify_file


def detectar_variables_en_escrito(docx_path):
    """
    Detecta todas las variables en el documento .docx que siguen el formato {{variable}}.
    """
    doc = Document(docx_path)
    texto_completo = "".join([para.text + "\n" for para in doc.paragraphs])
    pattern = r"\{\{(.*?)\}\}"
    return re.findall(pattern, texto_completo)


def obtener_valores_desde_excel(row, variables):
    """
    Obtiene los valores de las variables desde una fila específica de un DataFrame.
    """
    return {variable: row.get(variable, None) for variable in variables}


def reemplazar_variables_en_escrito(texto, valores_variables):
    """
    Reemplaza las variables en el texto del documento con los valores proporcionados.
    """
    for variable, valor in valores_variables.items():
        if valor is not None:
            pattern = r"\{\{\s*" + re.escape(variable) + r"\s*\}\}"
            texto = re.sub(pattern, str(valor), texto)
    return texto


def leer_documento_como_texto(docx_path):
    """
    Lee el contenido del archivo .docx como texto.
    """
    doc = Document(docx_path)
    return "".join([para.text + "\n" for para in doc.paragraphs])


def execute_steps(driver, texto_modificado, id_sac, user_config):
    """
    Ejecuta los steps necesarios utilizando Selenium.
    """
    user_login(driver, user_config)
    open_new_tab(driver)
    navigate_and_select(driver)
    select_document_type(driver, texto_modificado)
    input_case_number(driver, id_sac)
    skip_attachments(driver)
    pre_submit(driver)
    submit_case(driver)


def main(excel_path, docx_path, user_key):
    # Verificar que el archivo Word exista
    if not verify_file(docx_path):
        print("El archivo .docx no se encontró. Finalizando script.")
        return

    # Obtener configuración del usuario
    user_config = users.get(user_key.upper())
    if not user_config:
        print(f"Usuario {user_key} no encontrado en la configuración.")
        return

    # Detectar variables en el documento Word
    variables_encontradas = detectar_variables_en_escrito(docx_path)
    print(f"Variables encontradas en el documento: {variables_encontradas}")

    # Leer archivo Excel
    try:
        df = pd.read_excel(excel_path, engine='openpyxl')
    except Exception as e:
        print(f"Error al leer el archivo Excel: {e}")
        return

    if df.empty:
        print("El archivo Excel está vacío. Finalizando script.")
        return

    # Crear o verificar la columna Observación
    if "Observación" not in df.columns:
        df["Observación"] = None
    print("Columna 'Observación' creada.")

    for index, row in df.iterrows():
        # Saltar si el escrito ya fue presentado
        if row.get("Observación") == "Escrito presentado":
            print(f"Fila {index}: El título ya fue presentado. Saltando.")
            continue

        id_sac = row.get("IdSAC")
        if pd.isnull(id_sac):
            print(f"Fila {index}: IdSAC ausente. Saltando.")
            df.at[index, "Observación"] = "IdSAC ausente"
            df.to_excel(excel_path, index=False)
            continue

        # Obtener valores de las variables
        valores_variables = obtener_valores_desde_excel(row, variables_encontradas)

        # Leer y modificar el texto del documento Word
        texto_completo = leer_documento_como_texto(docx_path)
        texto_modificado = reemplazar_variables_en_escrito(texto_completo, valores_variables)

        # Configurar Selenium WebDriver
        driver = setup_driver()
        try:
            execute_steps(driver, texto_modificado, id_sac, user_config)
            print(f"Expediente con IdSAC {id_sac} procesado exitosamente.")
            df.at[index, "Observación"] = "Escrito presentado"
        except Exception as e:
            print(f"Ocurrió un error con el expediente {id_sac}: {e}")
            df.at[index, "Observación"] = f"Error: {str(e)}"
        finally:
            driver.quit()

        # Guardar el archivo después de cada fila procesada
        df.to_excel(excel_path, index=False, engine='openpyxl')
        print(f"Archivo actualizado guardado en: {excel_path}")

    print("Procesamiento completado.")


if __name__ == "__main__":
    base_path = os.path.dirname(os.path.abspath(__file__))
    excel_path = os.path.join(base_path, "tables", "SONZINI CAUSAS PARA PEDIDO DE ARCHIVO.xlsx")
    docx_path = os.path.join(base_path, "certificates", "pedido archivo de causa.docx")
    user_key = "SONZINI"
    main(excel_path, docx_path, user_key)
