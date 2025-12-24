import os
import pandas as pd
import re
from docx import Document

from config import users
from steps.user_login import user_login
from steps.manage_tabs import open_new_tab
from steps.driver_setup import setup_driver
from steps.navigate_and_select import navigate_and_select
from steps.select_document_type import select_document_type
from steps.input_case_number import input_case_number
from steps.upload_one_attachment import upload_one_attachment
from steps.submit_case import save_without_submission, submit_case
from steps.pre_submit import pre_submit
from steps.verify_file import verify_file


def detectar_variables_en_escrito(docx_path):
    """
    Detecta todas las variables en el documento .docx que siguen el formato {{variable}}.
    """
    doc = Document(docx_path)
    texto_completo = "".join([para.text + "\n" for para in doc.paragraphs])
    pattern = r"\{\{(.*?)\}\}"
    return re.findall(pattern, texto_completo)


def obtener_valores_desde_excel(row, variables):
    """
    Obtiene los valores de las variables desde una fila específica de un DataFrame.
    """
    return {variable: row.get(variable, None) for variable in variables}


def reemplazar_variables_en_escrito(texto, valores_variables):
    """
    Reemplaza las variables en el texto del documento con los valores proporcionados.
    """
    for variable, valor in valores_variables.items():
        if valor is not None:
            pattern = r"\{\{\s*" + re.escape(variable) + r"\s*\}\}"
            texto = re.sub(pattern, str(valor), texto)
    return texto


def leer_documento_como_texto(docx_path):
    """
    Lee el contenido del archivo .docx como texto.
    """
    doc = Document(docx_path)
    return "".join([para.text + "\n" for para in doc.paragraphs])


def execute_steps(driver, texto_modificado, id_sac, user_config, attachment_path, document_type_text):
    """
    Ejecuta los steps necesarios utilizando Selenium, subiendo solo un archivo adjunto.
    """
    user_login(driver, user_config)
    open_new_tab(driver)
    navigate_and_select(driver)
    select_document_type(driver, texto_modificado, document_type_text)
    input_case_number(driver, id_sac)
    upload_one_attachment(driver, attachment_path)
    pre_submit(driver)
    submit_case(driver)



if __name__ == "__main__":
    base_path = os.path.dirname(os.path.abspath(__file__))
    excel_path = os.path.join(base_path, "tables", "Pedir CNOE (19-12-2025).xlsx")
    docx_path = os.path.join(base_path, "certificates", "CERTIFICADO NO OPOSICION DE EXCEPCIONES-Solicita.docx")
    user_key = "FONSECA"
    attachments_folder = "Pedir CNOE (19-12-2025)"
    attachments_base_path = os.path.join(os.path.dirname(base_path), "titles")
    col_idsac = "IdSAC"
    col_archivo = "Archivo"
    col_observacion = "Observación"
    observacion_ok = "Escrito presentado"
    document_type_text = "CERTIFICACION DE NO OPOSICIÓN DE EXCEPCIONES - SOLICITA"

    # Verificación de archivos básicos
    if not verify_file(docx_path):
        print("El archivo .docx no se encontró. Finalizando script.")
        exit(1)

    user_config = users.get(user_key.upper())
    if not user_config:
        print(f"Usuario {user_key} no encontrado en la configuración.")
        exit(1)

    variables_encontradas = detectar_variables_en_escrito(docx_path)
    print(f"Variables encontradas en el documento: {variables_encontradas}")

    df = pd.read_excel(excel_path)
    if df.empty:
        print("El archivo Excel está vacío. Finalizando script.")
        exit(1)

    if col_observacion not in df.columns:
        df[col_observacion] = None
    print(f"Columna '{col_observacion}' disponible.")

    for index, row in df.iterrows():
        if row.get(col_observacion) == observacion_ok:
            print(f"Fila {index}: El escrito ya fue presentado anteriormente. Saltando esta fila.")
            continue

        id_sac = row.get(col_idsac)
        if pd.isnull(id_sac):
            print(f"Fila {index}: IdSAC ausente. Se marcará en Observación, pero seguimos.")
            df.at[index, col_observacion] = "IdSAC ausente"

        archivo = row.get(col_archivo)
        if pd.isnull(archivo):
            print(f"Fila {index}: Archivo ausente. Se marcará en Observación, pero seguimos.")
            df.at[index, col_observacion] = "Archivo ausente"
            continue

        attachment_path = os.path.join(attachments_base_path, attachments_folder, str(archivo))
        print(f"Procesando fila {index} con IdSAC {id_sac} y archivo '{archivo}'...")

        texto_completo = leer_documento_como_texto(docx_path)
        valores = obtener_valores_desde_excel(row, variables_encontradas)
        texto_modificado = reemplazar_variables_en_escrito(texto_completo, valores)

        driver = setup_driver()
        try:
            execute_steps(
                driver,
                texto_modificado,
                id_sac,
                user_config,
                attachment_path,
                document_type_text
            )
            df.at[index, col_observacion] = observacion_ok
            print(f"Fila {index}: Escrito presentado correctamente.")
        except Exception as e:
            print(f"Error en fila {index}, IdSAC {id_sac}: {e}")
            df.at[index, col_observacion] = f"Error: {e}"
        finally:
            driver.quit()
            df.to_excel(excel_path, index=False)
            print(f"Progreso guardado después de fila {index}.")

    print("Procesamiento de todas las filas completado.")
