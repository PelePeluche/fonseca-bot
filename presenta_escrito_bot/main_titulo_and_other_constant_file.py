import os
import pandas as pd
import re
from docx import Document
from config import users
from steps.user_login import user_login
from steps.manage_tabs import open_new_tab
from steps.driver_setup import setup_driver
from steps.navigate_and_select import navigate_and_select
from steps.select_document_type import select_document_type
from steps.input_case_number import input_case_number
from steps.upload_titulo_aporte_poder import upload_titulo_aporte_poder
from steps.submit_case import save_without_submission, submit_case
from steps.pre_submit import pre_submit
from steps.verify_file import verify_file
from config import table_name, titles_foldername, planillas_foldername

def detectar_variables_en_escrito(docx_path):
    """
    Detecta todas las variables en el documento .docx que siguen el formato {{variable}}.
    """
    doc = Document(docx_path)
    texto_completo = "".join([para.text + "\n" for para in doc.paragraphs])
    pattern = r"\{\{(.*?)\}\}"
    return re.findall(pattern, texto_completo)

def obtener_valores_desde_excel(row, variables):
    """
    Obtiene los valores de las variables desde la fila `row`.
    Si la columna no existe o es NaN, retorna cadena vacía.
    """
    result = {}
    for var in variables:
        val = row.get(var, "")
        # Si val es NaN, lo forzamos a ""
        if pd.isnull(val):
            val = ""
        result[var] = val
    return result

def reemplazar_variables_en_escrito(texto, valores_variables):
    """
    Reemplaza las variables en el texto del documento con los valores proporcionados.
    """
    for variable, valor in valores_variables.items():
        # Como ahora valor nunca será None, no es necesario condicional.
        pattern = r"\{\{\s*" + re.escape(variable) + r"\s*\}\}"
        texto = re.sub(pattern, str(valor), texto)
    return texto

def leer_documento_como_texto(docx_path):
    """
    Lee el contenido del archivo .docx como texto.
    """
    doc = Document(docx_path)
    return "".join([para.text + "\n" for para in doc.paragraphs])

def execute_steps(driver, texto_modificado, id_sac, user_config, title_path, planilla_path):
    """
    Ejecuta los steps necesarios utilizando Selenium.
    """
    user_login(driver, user_config)
    open_new_tab(driver)
    navigate_and_select(driver)
    select_document_type(driver, texto_modificado)
    input_case_number(driver, id_sac)
    upload_titulo_aporte_poder(driver, title_path, planilla_path)
    pre_submit(driver)
    #save_without_submission(driver)
    submit_case(driver)

def main(excel_path, docx_path, user_key):
    # Verificamos que el .docx existe
    if not verify_file(docx_path):
        print("El archivo .docx no se encontró. Finalizando script.")
        return

    # Cargamos la configuración del usuario
    user_config = users.get(user_key.upper())
    if not user_config:
        print(f"Usuario {user_key} no encontrado en la configuración.")
        return

    # Detectamos variables en el documento
    variables_encontradas = detectar_variables_en_escrito(docx_path)
    print(f"Variables encontradas en el documento: {variables_encontradas}")

    # Leemos el archivo Excel
    df = pd.read_excel(excel_path)
    if df.empty:
        print("El archivo Excel está vacío. Finalizando script.")
        return

    # Creamos la columna Observación si no existe
    if "Observación" not in df.columns:
        df["Observación"] = None
    print("Columna 'Observación' creada.")

    for index, row in df.iterrows():
        # Verificamos si el título ya fue rectificado en una corrida anterior
        if row.get("Observación") == "Escrito presentado":
            print(f"Fila {index}: El título ya fue rectificado anteriormente. Saltando esta fila.")
            continue

        # Obtenemos el IdSAC
        id_sac = row.get("IdSAC")
        if pd.isnull(id_sac):
            print(f"Fila {index}: IdSAC ausente. Se marcará en Observación, pero seguimos.")
            df.at[index, "Observación"] = "IdSAC ausente"

        # Obtenemos valores de las variables desde la fila
        valores_variables = obtener_valores_desde_excel(row, variables_encontradas)

        # Leer y modificar el texto del documento
        texto_completo = leer_documento_como_texto(docx_path)
        texto_modificado = reemplazar_variables_en_escrito(
            texto_completo, valores_variables
        )

        # Tomamos el valor de la columna "Archivo"
        archivo = row.get("Archivo")
        title_path = None
        if not pd.isnull(archivo):
            # Si no es nulo, construimos la ruta
            archivo_str = str(archivo)
            title_path = os.path.join(
                os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                'titles',
                titles_foldername,
                archivo_str,
            )

        print(f"Procesando expediente con IdSAC {id_sac}...")
        print(f"Path para el archivo: {title_path if title_path else 'None'}")

        # Usar siempre la misma planilla fija
        planilla_fija = "AFIP - FONSECA 2025.pdf"
        planilla_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            'constants_files',
            planilla_fija,
        )
        print(f"Path para la planilla: {planilla_path}")

        # Iniciamos el driver Selenium
        driver = setup_driver()
        try:
            execute_steps(
                driver, texto_modificado, id_sac, user_config, title_path, planilla_path
            )
            print(f"Expediente con IdSAC {id_sac} procesado exitosamente.")
            df.at[index, "Observación"] = "Escrito presentado"
        except Exception as e:
            # print(f"Ocurrió un error con el expediente {id_sac}: {e}")
            # df.at[index, "Observación"] = f"Error: {str(e)}"
            continue
        finally:
            driver.quit()

        # Guardamos los cambios en el Excel tras cada fila
        df.to_excel(excel_path, index=False)
        print(f"Archivo actualizado guardado en: {excel_path}")

    print("Procesamiento completado.")


if __name__ == "__main__":
    base_path = os.path.dirname(os.path.abspath(__file__))
    excel_path = os.path.join(base_path, "tables", table_name)
    docx_path = os.path.join(base_path, "certificates", "MODELO_ EJECUCIÓN - PROCEDIMIENTO DE EJECUCIÓN FISCAL ADMINISTRATIVA.docx")
    user_key = "FONSECA"
    main(excel_path, docx_path, user_key)
